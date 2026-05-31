from __future__ import annotations

import argparse
import json
import os
import re
import ssl
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import certifi
import imageio_ffmpeg
import whisper
import yt_dlp
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from youtube_transcript_api import YouTubeTranscriptApi

import os
import imageio_ffmpeg
from pathlib import Path

PROJECT_DIR = Path(__file__).resolve().parent
LOCAL_BIN = PROJECT_DIR / "bin"
os.environ["PATH"] = str(LOCAL_BIN) + os.pathsep + os.environ.get("PATH", "")

ffmpeg_bin = Path(imageio_ffmpeg.get_ffmpeg_exe()).parent

os.environ["PATH"] = (
    str(ffmpeg_bin)
    + os.pathsep
    + os.environ.get("PATH", "")
)

ssl._create_default_https_context = lambda: ssl.create_default_context(cafile=certifi.where())

VIDEO_SECTION_HEADERS = {
    "Synthesis (Q1)": "Synthesis",
    "Rhetorical Analysis (Q2)": "Rhetorical Analysis",
    "Argument (Q3)": "Argument",
    "General Tips on Commentary": "Commentary",
    "Sophistication": "Sophistication",
}


@dataclass
class VideoResult:
    title: str
    category: str
    url: Optional[str] = None
    video_id: Optional[str] = None
    duration_seconds: Optional[int] = None
    transcript: Optional[str] = None
    transcript_available: bool = False
    transcript_source: Optional[str] = None
    notes: Optional[Dict[str, object]] = None
    error: Optional[str] = None
    manual_review: bool = False


def read_pdf_text(pdf_path: str | Path) -> str:
    reader = PdfReader(str(pdf_path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_video_titles(assignment_text: str) -> List[VideoResult]:
    lines = [line.strip() for line in assignment_text.splitlines() if line.strip()]
    current_category = "Priority"
    videos: List[VideoResult] = []

    for line in lines:
        clean = re.sub(r"^●\s*", "", line).strip()

        for header, category in VIDEO_SECTION_HEADERS.items():
            if header.lower() in clean.lower():
                current_category = category

        if line.startswith("●"):
            title = clean
            if title.lower() in {"rhetorical analysis", "synthesis", "argument"}:
                continue
            videos.append(VideoResult(title=title, category=current_category))

    seen = set()
    unique = []
    for video in videos:
        key = video.title.lower()
        if key not in seen:
            seen.add(key)
            unique.append(video)

    return unique


def search_youtube(title: str) -> Tuple[Optional[str], Optional[str], Optional[int], Optional[str]]:
    options = {
        "quiet": True,
        "skip_download": True,
        "extract_flat": False,
        "default_search": "ytsearch1",
        "noplaylist": True,
    }

    try:
        with yt_dlp.YoutubeDL(options) as ydl:
            info = ydl.extract_info(f"ytsearch1:{title}", download=False)
            entries = info.get("entries") or []
            if not entries:
                return None, None, None, "No YouTube result found."

            video = entries[0]
            return video.get("webpage_url"), video.get("id"), video.get("duration"), None

    except Exception as exc:
        return None, None, None, str(exc)


def fetch_transcript(video_id: str) -> Tuple[Optional[str], Optional[str]]:
    try:
        ytt_api = YouTubeTranscriptApi()

        if hasattr(YouTubeTranscriptApi, "get_transcript"):
            transcript_items = YouTubeTranscriptApi.get_transcript(video_id)
        elif hasattr(ytt_api, "fetch"):
            transcript_obj = ytt_api.fetch(video_id)
            transcript_items = transcript_obj.to_raw_data()
        else:
            return None, "No usable transcript method found."

        transcript = " ".join(item.get("text", "") for item in transcript_items)
        transcript = re.sub(r"\s+", " ", transcript).strip()

        if not transcript:
            return None, "Transcript was empty."

        return transcript, None

    except Exception as exc:
        return None, f"Transcript unavailable: {exc}"


def transcribe_with_whisper(video_url: str, video_id: str) -> Tuple[Optional[str], Optional[str]]:
    try:
        audio_dir = Path("audio")
        audio_dir.mkdir(exist_ok=True)

        output_template = str(audio_dir / f"{video_id}.%(ext)s")
        ffmpeg_path = imageio_ffmpeg.get_ffmpeg_exe()
        os.environ["PATH"] = str(Path(ffmpeg_path).parent) + os.pathsep + os.environ.get("PATH", "")

        subprocess.run(
            [
                "yt-dlp",
                "-x",
                "--audio-format",
                "mp3",
                "--ffmpeg-location",
                ffmpeg_path,
                "-o",
                output_template,
                video_url,
            ],
            check=True,
        )

        audio_path = audio_dir / f"{video_id}.mp3"

        if not audio_path.exists():
            candidates = list(audio_dir.glob(f"{video_id}.*"))
            if not candidates:
                return None, "Whisper fallback failed: no audio file downloaded."
            audio_path = candidates[0]

        print(f"    Transcribing {audio_path}")

        model = whisper.load_model("base")
        result = model.transcribe(str(audio_path))
        transcript = result.get("text", "").strip()

        if not transcript:
            return None, "Whisper fallback failed: transcript was empty."

        print(f"    Transcript length: {len(transcript)} characters")

        return transcript, None

    except Exception as exc:
        return None, f"Whisper fallback failed: {exc}"


def duration_label(seconds: Optional[int]) -> str:
    if seconds is None:
        return "Unknown"
    return f"{seconds // 60}:{seconds % 60:02d}"


def load_agent_instructions() -> str:
    prompt_path = Path("prompts/worksheet_prompt.txt")
    if prompt_path.exists():
        content = prompt_path.read_text(encoding="utf-8").strip()
        if content:
            return content

    return """
You are an AP English Language worksheet agent.
Complete one worksheet page for each video.
Use specific details from the transcript to prove the video was watched.
Do not invent details.
Write in a thoughtful high school student voice.
""".strip()


def fallback_notes(video: VideoResult) -> Dict[str, object]:
    return {
        "before_video": (
            f"I chose this video because the title, \"{video.title},\" points to a skill that matters for AP Lang FRQs. "
            "I want to understand the process well enough to use it under timed exam conditions instead of relying on vague advice."
        ),
        "lesson_overview": (
            "Manual review needed. The agent could not produce a reliable transcript-based summary for this video. "
            "After watching it directly, this paragraph should include the speaker's main advice, a concrete example or phrase from the video, and an explanation of how the advice applies to an AP Lang FRQ."
        ),
        "learned": [
            "Manual review needed: add one concrete strategy from the video.",
            "Manual review needed: add one specific example, phrase, acronym, or distinction from the speaker.",
            "Manual review needed: add one mistake the video warns students to avoid.",
        ],
        "already_knew": [
            "I already knew that AP Lang FRQs need a defensible thesis rather than a broad topic sentence.",
            "I already knew that evidence has to be explained instead of simply dropped into a paragraph.",
        ],
        "exam_day": "Use one specific piece of advice from the video when revising this entry before submission.",
    }


def generate_notes_with_openai(video: VideoResult, client: OpenAI, model: str) -> Dict[str, object]:
    transcript_excerpt = (video.transcript or "")[:18000]
    instructions = load_agent_instructions()

    prompt = f"""
{instructions}

Video title: {video.title}
Video category: {video.category}
Video length: {duration_label(video.duration_seconds)}
Transcript source: {video.transcript_source}

Transcript excerpt:
{transcript_excerpt}

Return valid JSON only with this exact shape:
{{
  "before_video": "2 to 4 complete sentences about what skill the student hopes to understand before watching.",
  "lesson_overview": "One substantial paragraph with transcript-specific details, examples, phrases, frameworks, or distinctions.",
  "learned": ["specific learned item 1", "specific learned item 2", "specific learned item 3"],
  "already_knew": ["plausible thing the student already knew 1", "plausible thing the student already knew 2"],
  "exam_day": "one practical sentence about what to remember on exam day"
}}

Rules:
- Do not invent details that are not in the transcript.
- Be specific enough to prove the video was watched.
- Avoid vague filler.
""".strip()

    response = client.responses.create(
        model=model,
        input=prompt,
        temperature=0.35,
    )

    raw = response.output_text.strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", raw, flags=re.S)
        if not match:
            raise ValueError(f"OpenAI did not return JSON. Output was: {raw[:500]}")
        return json.loads(match.group(0))


def process_all_videos(videos: List[VideoResult], model_name: str) -> List[VideoResult]:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    client = OpenAI(api_key=api_key) if api_key else None

    for index, video in enumerate(videos, start=1):
        print(f"[{index}/{len(videos)}] Processing: {video.title}")

        url, video_id, duration, search_error = search_youtube(video.title)
        video.url = url
        video.video_id = video_id
        video.duration_seconds = duration

        if search_error:
            video.error = search_error
            video.manual_review = True
            video.notes = fallback_notes(video)
            continue

        if video.video_id:
            transcript, transcript_error = fetch_transcript(video.video_id)

            if transcript:
                video.transcript = transcript
                video.transcript_available = True
                video.transcript_source = "youtube_transcript_api"
                video.manual_review = False
                video.error = None
            elif video.url:
                print("    Caption transcript failed. Trying Whisper fallback...")
                transcript, whisper_error = transcribe_with_whisper(video.url, video.video_id)

                if transcript:
                    video.transcript = transcript
                    video.transcript_available = True
                    video.transcript_source = "whisper_audio_fallback"
                    video.manual_review = False
                    video.error = None
                else:
                    video.transcript_available = False
                    video.manual_review = True
                    video.error = whisper_error or transcript_error
            else:
                video.error = transcript_error
                video.manual_review = True
        else:
            video.error = "No video ID found."
            video.manual_review = True

        if client and video.transcript_available:
            try:
                video.notes = generate_notes_with_openai(video, client, model_name)
                video.manual_review = False
            except Exception as exc:
                video.error = f"OpenAI note generation failed: {exc}"
                video.manual_review = True
                video.notes = fallback_notes(video)
        else:
            if not client:
                video.error = "No OPENAI_API_KEY found in .env."
            video.manual_review = True
            video.notes = fallback_notes(video)

    return videos


def clean_docx_text(value) -> str:
    if value is None:
        return ""
    text = str(value)
    return "".join(ch for ch in text if ch == "\n" or ch == "\t" or ord(ch) >= 32)


def add_labeled_paragraph(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    p.add_run(clean_docx_text(label)).bold = True
    p.add_run(clean_docx_text(body))


def export_docx(videos: List[VideoResult], output_path: str | Path) -> None:
    doc = Document()
    doc.add_heading("Choose Your Own Adventure FRQ Exam Review Notes", level=1)
    doc.add_paragraph("All listed videos processed. Each video has its own worksheet page.")

    for number, video in enumerate(videos, start=1):
        if number > 1:
            doc.add_page_break()

        notes = video.notes or fallback_notes(video)

        doc.add_heading(clean_docx_text(f"Video {number}: {video.title}"), level=2)
        add_labeled_paragraph(doc, "Category: ", video.category)
        add_labeled_paragraph(doc, "Video Link: ", video.url or "Manual link needed")
        add_labeled_paragraph(doc, "Video Length: ", duration_label(video.duration_seconds))
        add_labeled_paragraph(doc, "Transcript Source: ", video.transcript_source or "None")

        add_labeled_paragraph(
            doc,
            "BEFORE THE VIDEO: Based on the video description, what skill are you hoping this video will help you better understand?\n",
            str(notes.get("before_video", "")),
        )

        add_labeled_paragraph(doc, "Lesson Overview:\n", str(notes.get("lesson_overview", "")))

        doc.add_paragraph("3 Things I Learned").runs[0].bold = True
        for i, item in enumerate(notes.get("learned", [])[:3], start=1):
            doc.add_paragraph(clean_docx_text(f"{i}. {item}"))

        doc.add_paragraph("2 Things I Already Knew").runs[0].bold = True
        for i, item in enumerate(notes.get("already_knew", [])[:2], start=1):
            doc.add_paragraph(clean_docx_text(f"{i}. {item}"))

        add_labeled_paragraph(
            doc,
            "1 Thing I Want to Remember on Exam Day Related to the Video:\n",
            str(notes.get("exam_day", "")),
        )

        if video.manual_review:
            add_labeled_paragraph(doc, "Manual Review Flag: ", video.error or "Needs manual review.")

    doc.save(str(output_path))


def export_report(videos: List[VideoResult], report_path: str | Path) -> None:
    rows = []
    for video in videos:
        rows.append(
            {
                "title": video.title,
                "category": video.category,
                "url": video.url,
                "video_id": video.video_id,
                "duration": duration_label(video.duration_seconds),
                "transcript_available": video.transcript_available,
                "transcript_source": video.transcript_source,
                "manual_review": video.manual_review,
                "error": video.error,
            }
        )

    Path(report_path).write_text(json.dumps(rows, indent=2), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="AP Lang all-videos worksheet agent")
    parser.add_argument("--assignment", required=True, help="Path to assignment PDF")
    parser.add_argument("--output", default="completed_all_ap_lang_frq_video_notes.docx")
    parser.add_argument("--report", default="video_processing_report.json")
    parser.add_argument("--model", default="gpt-4.1-mini")
    args = parser.parse_args()

    assignment_text = read_pdf_text(args.assignment)
    videos = extract_video_titles(assignment_text)

    if not videos:
        raise RuntimeError("No video titles found in the assignment PDF.")

    print(f"Found {len(videos)} unique videos.")

    processed = process_all_videos(videos, args.model)

    export_docx(processed, args.output)
    export_report(processed, args.report)

    manual_count = sum(1 for video in processed if video.manual_review)
    transcript_count = sum(1 for video in processed if video.transcript_available)

    print(f"Done. Wrote {args.output}")
    print(f"Report saved to {args.report}")
    print(f"Transcript entries: {transcript_count}")
    print(f"Manual-review entries: {manual_count}")


    google_doc_link = upload_docx_to_google_drive(
        args.output,
        "Completed AP Lang FRQ Video Notes",
    )

    if google_doc_link:
        print(f"Google Doc created: {google_doc_link}")


def upload_docx_to_google_drive(docx_path: str, google_doc_name: str) -> Optional[str]:
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    scopes = ["https://www.googleapis.com/auth/drive.file"]
    creds = None

    token_path = Path("token.json")
    credentials_path = Path("credentials.json")

    if token_path.exists():
        creds = Credentials.from_authorized_user_file(str(token_path), scopes)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not credentials_path.exists():
                print("Google upload skipped: credentials.json not found.")
                return None

            flow = InstalledAppFlow.from_client_secrets_file(
                str(credentials_path),
                scopes,
            )
            creds = flow.run_local_server(port=0)

        token_path.write_text(creds.to_json(), encoding="utf-8")

    service = build("drive", "v3", credentials=creds)

    file_metadata = {
        "name": google_doc_name,
        "mimeType": "application/vnd.google-apps.document",
    }

    media = MediaFileUpload(
        docx_path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=True,
    )

    uploaded = (
        service.files()
        .create(
            body=file_metadata,
            media_body=media,
            fields="id, webViewLink",
        )
        .execute()
    )

    return uploaded.get("webViewLink")


if __name__ == "__main__":
    main()
